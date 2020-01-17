# -*- coding: utf-8  -*-
import io
import os
import sys

from docx import Document
from docx.shared import Inches, Cm
import cv2
import numpy as np

scale = 37.79527559055118  # word的cm与px的比例


def main(param):
    docx_name, chessboard_cm, width_cm, height_cm = param  		# 获取参数
    chessboard_pixel = chessboard_cm*scale						# 棋盘格单元对应像素
    width = round(width_cm / chessboard_cm)						# 棋盘格角点列数
    height = round(height_cm / chessboard_cm)					# 棋盘格角点行数
    width_pix = round(width_cm * scale)							# 棋盘格列像素
    height_pix = round(height_cm * scale)						# 棋盘格行像素

    size = width_cm/height_cm									# 棋盘格长宽比例(用于最后的裁切)

    image = np.zeros((height_pix, width_pix, 3), dtype=np.uint8)# 生成对应像素图片
    image.fill(255)												# 填充白色
    color = (255, 255, 255)
    fill_color = 0
    for j in range(0, height + 1):								# 绘制棋盘
        y = round(j * chessboard_pixel)
        for i in range(0, width + 1):
            x0 = round(i * chessboard_pixel)
            y0 = y
            rect_start = (x0, y0)

            x1 = round(x0 + chessboard_pixel)
            y1 = round(y0 + chessboard_pixel)
            rect_end = (x1, y1)
            cv2.rectangle(image, rect_start, rect_end, color, 1, 0)
            image[y0:y1, x0:x1] = fill_color
            if width % 2:
                if i != width:
                    fill_color = (0 if (fill_color == 255) else 255)
            else:
                if i != width + 1:
                    fill_color = (0 if (fill_color == 255) else 255)
    bottom = round(width_pix/size)								# 获取按cm比例对应的行像素
    if bottom < height_pix:
        image = image[0:bottom, :, :]							# 裁切

    # 创建显示窗口
    win_name = "chessboard"
    # cv.namedWindow(win_name, cv.WINDOW_NORMAL)
    cv2.imwrite(win_name + ".bmp", image)
    # cv.imshow(win_name, image)
    # cv.waitKey()

    doc = Document()  											# 以默认模板建立文档对象
    distance = Inches(0)
    sec = doc.sections[0]										# 页边距0英寸
    sec.left_margin = distance  								# 以下依次设置左、右、上、下页面边距
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance

    sec.page_width = Cm(width_cm)  								# 设置页面宽度
    sec.page_height = Cm(height_cm)  							# 设置页面高度

    img_encode = cv2.imencode('.bmp', image)[1]					# 对opencv图片进行编码
    str_encode = img_encode .tostring()
    cc = io.BytesIO(str_encode)
    # img = doc.add_picture(cc, Cm(42.01))
    doc.add_picture(cc)											# 插入图片至word
    # doc.add_picture(win_name + ".bmp")
    doc.save(docx_name)       									# 保存图像


if __name__ == '__main__':
    param = ['chessboard.docx', 0.7, 21, 29.7]  # 文档名, 正方格长度cm, 页面宽度cm, 页面高度cm
    length = len(sys.argv)
    if length > 1:
        for idx in range(1, 5 if (length > 5) else length):
            param[idx - 1] = sys.argv[idx]
        strlist = os.path.splitext(str(param[0]))
        if len(strlist) == 1:
            param[0] = str(param[0]) + '.docx'
        elif strlist[-1] != 'docx':
            param[0] = ''.join(strlist[0:-1]) + '.docx'
        # print(param)
    # if param[1]*scale > 1:
    #     print('生成棋盘格失败, 单元格边长过小, 像素解析力不足, 无法生成!')
    # else:
    main(param)
    print('生成棋盘格成功. 文件名:%s, 单元格边长:%0.2fcm, 棋盘格宽度:%0.2fcm, 棋盘格高度:%0.2fcm'
          % (param[0], param[1], param[2], param[3]))
